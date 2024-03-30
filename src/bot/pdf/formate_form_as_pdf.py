import os
from docx import Document

from reportlab.pdfgen import canvas
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_RIGHT, TA_LEFT
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

styles = getSampleStyleSheet() # дефолтовые стили
# the magic is here
styles['Normal'].fontName='TimesNewRoman'
styles['Heading1'].fontName='TimesNewRoman'

pdfmetrics.registerFont(TTFont('TimesNewRoman','TimesNewRoman.ttf', 'UTF-8'))


# def hello(c):
#     c.drawString(100,100,'Hello World')
# c = canvas.Canvas(f"{application}.pdf'"')
# hello(c)
# c.showPage()
# c.save()


style_right = ParagraphStyle(
        name='left',
        parent=styles['Normal'],
        fontSize=12,
        fontName='TimesNewRoman',
        spaceBefore=3,
        spaceAfter=3,
        alignment=TA_RIGHT)

style_left = ParagraphStyle(
        name='left',
        parent=styles['Normal'],
        fontSize=12,
        fontName='TimesNewRoman',
        alignment=TA_LEFT)

style_center = ParagraphStyle(
        name='left',
        parent=styles['Normal'],
        fontSize=12,
        fontName='TimesNewRoman',
        alignment=TA_CENTER)


#project_dir = os.path.dirname(__file__)
file_path = (r'C:\Users\Хозяин\Desktop\24.03.2024\src\bot\common\form_template.docx')
#file_path = os.path.join(project_dir, '\common', 'form_template.docx')

def formate_form():
    doc = []
    docx = Document(file_path)
    zayava_flag = False

    for line in docx.paragraphs:

        if zayava_flag:
            style=style_left
        elif 'Заява' == line.text:
            style=style_center
            zayava_flag = True
        else:
            style=style_right

        print(line.text)
        if line.text == '':
            doc.append(Spacer(1,12))
        doc.append(Paragraph(line.text, style))
    return doc

# def convert_docx_to_txt(docx_file, txt_file):
#     doc = Document(docx_file)
#     with open(txt_file, 'w', encoding='utf-8') as txt:
#         for paragraph in doc.paragraphs:
#             txt.write(paragraph.text + '\n')

# # Пример использования
# txt_file = 'example.txt'
# convert_docx_to_txt(file_path, txt_file)

doc = formate_form()

SimpleDocTemplate(f'form.pdf',path='',
                  rightMagrin=2.54*cm, leftMagrin=2.54*cm,
                  topMargin=2.54*cm, bottomMagrin=2.54*cm).build(doc)